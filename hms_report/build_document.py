"""Build the HMS report .docx document (month, quarter, or year)."""

from __future__ import annotations

import io
from datetime import datetime

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.table import Table

from .config import get_report_title, get_report_type_label

FONT_NAME = "Calibri"
HEADER_COLOR = RGBColor(0x1F, 0x38, 0x64)
HEADER_TEXT_COLOR = RGBColor(0xFF, 0xFF, 0xFF)
ALT_ROW_COLOR = RGBColor(0xF2, 0xF2, 0xF2)


def _set_cell_shading(cell, color: RGBColor) -> None:
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): f"{color}",
        qn("w:val"): "clear",
    })
    shading.append(shd)


def _header_cell(row, col: int, text: str) -> None:
    cell = row.cells[col]
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.color.rgb = HEADER_TEXT_COLOR
    _set_cell_shading(cell, HEADER_COLOR)


def _data_cell(
    row, col: int, text: str, *,
    center: bool = False, bold: bool = False,
    alt: bool = False, color: RGBColor | None = None,
    font_size: int = 8,
) -> None:
    cell = row.cells[col]
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    if alt:
        _set_cell_shading(cell, ALT_ROW_COLOR)


def _format_date(date_str: str | None) -> str:
    if not date_str:
        return ""
    try:
        d = datetime.fromisoformat(date_str.replace("Z", "+00:00"))
        return d.strftime("%d.%m.%Y")
    except (ValueError, TypeError):
        return str(date_str)[:10]


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = FONT_NAME
        run.font.color.rgb = HEADER_COLOR


def _add_body(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(11)
    run.font.bold = bold


def _add_chart(doc: Document, chart_bytes: bytes, width: float = 4.5) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(io.BytesIO(chart_bytes), width=Inches(width))


def _add_chart_row(doc: Document, buf1: bytes, buf2: bytes, buf3: bytes) -> None:
    """Add 3 pie charts side by side using a table."""
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Remove borders
    for cell in table.rows[0].cells:
        tc_pr = cell._element.get_or_add_tcPr()
        borders = tc_pr.makeelement(qn("w:tcBorders"), {})
        for side in ("top", "bottom", "left", "right"):
            border = borders.makeelement(qn(f"w:{side}"), {
                qn("w:val"): "none", qn("w:sz"): "0",
            })
            borders.append(border)
        tc_pr.append(borders)

    for i, buf in enumerate([buf1, buf2, buf3]):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(io.BytesIO(buf), width=Inches(2.0))


# ---- Section builders ----

def _build_cover(doc: Document, config: dict) -> None:
    report_type = config.get("reportType", "quarter")
    period_label = config.get("periodLabel", f"{config.get('quarter', '')} {config['year']}")

    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(get_report_title(report_type))
    run.font.name = FONT_NAME
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = HEADER_COLOR

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(period_label)
    run.font.name = FONT_NAME
    run.font.size = Pt(22)
    run.font.color.rgb = HEADER_COLOR

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    run = p.add_run("Aage Haverstad AS")
    run.font.name = FONT_NAME
    run.font.size = Pt(18)
    run.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Generert: {datetime.now().strftime('%d.%m.%Y')}")
    run.font.name = FONT_NAME
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()


def _build_toc(doc: Document, config: dict) -> None:
    report_type = config.get("reportType", "quarter")
    type_label = get_report_type_label(report_type)

    _add_heading(doc, "Innholdsfortegnelse")
    sections = [
        f"{type_label} HMS",
        "1. Kort status",
        "2. Hendelser og avvik",
        "3. Nokkelaktiviteter HMS",
        "4. Sykefravaer",
        "5. Tiltak neste periode",
        "6. Oversikt/Statistikk RUH:",
        "7. Oversikt/statistikk Kvalitetsavvik:",
        "8. Samlet oversikt HMSK:",
    ]
    for s in sections:
        p = doc.add_paragraph(s)
        for run in p.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(12)
    doc.add_page_break()


def _build_metadata(doc: Document, config: dict) -> None:
    report_type = config.get("reportType", "quarter")
    type_label = get_report_type_label(report_type)
    period_label = config.get("periodLabel", f"{config.get('quarter', '')} {config['year']}")

    _add_heading(doc, f"{type_label} HMS")
    _add_body(doc, "Aage Haverstad AS", bold=True)
    _add_body(doc, f"Periode: {period_label}")
    now = datetime.now()
    _add_body(doc, f"Dato: {now.strftime('%d. %B %Y')}")
    _add_body(doc, f"Utarbeidet av: {config['utarbeidetAv']}")


def _build_section1(doc: Document, config: dict) -> None:
    _add_heading(doc, "1. Kort status")
    _add_body(doc, config["manual"]["kortStatus"])


def _build_section2(doc: Document, data: dict) -> None:
    _add_heading(doc, "2. Hendelser og avvik")

    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"

    for i, h in enumerate(["Hendelse", "Antall", "Kommentar"]):
        _header_cell(table.rows[0], i, h)

    rows_data = [
        ("Skade/personulykke", "0", ""),
        ("Uoenskede hendelser (RUH)", str(data["rue"]["total"]), ""),
        ("Kvalitetsavvik (QD)", str(data["qd"]["total"]), ""),
    ]
    for i, (col1, col2, col3) in enumerate(rows_data):
        alt = i % 2 == 1
        _data_cell(table.rows[i + 1], 0, col1, alt=alt)
        _data_cell(table.rows[i + 1], 1, col2, center=True, alt=alt)
        _data_cell(table.rows[i + 1], 2, col3, alt=alt)

    closed = data["rue"]["byStatus"].get("Lukket", 0)
    open_rue = (
        data["rue"]["byStatus"].get("Aapen", 0)
        + data["rue"]["byStatus"].get("Ny", 0)
        + data["rue"]["byStatus"].get("Ubehandlet", 0)
    )
    _add_body(doc, f"Lukket avvik/RUH: {closed}")
    _add_body(doc, f"Aapne RUH: {open_rue}")
    _add_body(doc, f"Totalt: {data['rue']['total']} stk.")


def _build_section3(doc: Document, data: dict) -> None:
    _add_heading(doc, "3. Nokkelaktiviteter HMS")
    config = data["config"]
    nokkel = config["manual"]["nøkkelaktiviteter"]

    table = doc.add_table(rows=5, cols=3)
    table.style = "Table Grid"

    for i, h in enumerate(["Aktivitet", "Gjennomfoert", "Kommentar"]):
        _header_cell(table.rows[0], i, h)

    vr = data["vernerunder"]
    vr_comment = nokkel["vernerunder"]["kommentar"] or ", ".join(
        f"{m['label']}: {m['count']}" for m in vr["byMonth"]
    )
    sja = data["sja"]
    sja_comment = ", ".join(f"{m['label']}: {m['count']}" for m in sja["byMonth"])

    activities = [
        ("Vernerunder", f"{vr['total']} stk", vr_comment),
        ("HMS-moete", "Ja" if nokkel["hmsMøte"]["gjennomført"] else "Nei", nokkel["hmsMøte"]["kommentar"]),
        ("Risikovurdering", "Ja" if nokkel["risikovurdering"]["gjennomført"] else "Nei", nokkel["risikovurdering"]["kommentar"]),
        ("SJA (Sikker Jobb Analyse)", f"{sja['total']} stk", sja_comment),
    ]
    for i, (name, done, comment) in enumerate(activities):
        alt = i % 2 == 1
        _data_cell(table.rows[i + 1], 0, name, alt=alt)
        _data_cell(table.rows[i + 1], 1, done, center=True, alt=alt)
        _data_cell(table.rows[i + 1], 2, comment, alt=alt)


def _build_section4(doc: Document, config: dict) -> None:
    _add_heading(doc, "4. Sykefravaer")
    sf_list = config["manual"]["sykefravær"]
    start_month = int(config["dateRange"]["start"][5:7])
    end_month = int(config["dateRange"]["end"][5:7])
    months = list(range(start_month, end_month + 1))

    total_syk = 0
    total_egen = 0
    total_pct = 0.0
    count = 0
    for m in months:
        sf = sf_list[m - 1]
        total_syk += sf["sykmeldinger"]
        total_egen += sf["egenmeldinger"]
        total_pct += sf["fraværProsent"]
        count += 1

    avg = total_pct / count if count else 0
    _add_body(doc, f"Sykmeldinger: {total_syk}")
    _add_body(doc, f"Egenmeldinger: {total_egen}")
    _add_body(doc, f"Fravaersprosent: {avg:.1f}%")


def _build_section5(doc: Document, config: dict) -> None:
    _add_heading(doc, "5. Tiltak neste periode")
    for i, t in enumerate(config["manual"]["tiltakNesteKvartal"], 1):
        _add_body(doc, f"{i}. {t}")


def _build_rue_table(doc: Document, rows: list[dict]) -> None:
    headers = ["Hendelse", "Innsendt", "Prosjekt", "Tittel", "Nr.", "Type", "Omfattet", "Aarsak"]
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Table Grid"

    for i, h in enumerate(headers):
        _header_cell(table.rows[0], i, h)

    for i, r in enumerate(rows):
        alt = i % 2 == 1
        _data_cell(table.rows[i + 1], 0, _format_date(r.get("eventTime")), alt=alt)
        _data_cell(table.rows[i + 1], 1, _format_date(r["submitDate"]), alt=alt)
        _data_cell(table.rows[i + 1], 2, r["projectName"], alt=alt)
        _data_cell(table.rows[i + 1], 3, r["title"], alt=alt)
        _data_cell(table.rows[i + 1], 4, str(r["eventId"]), center=True, alt=alt)
        _data_cell(table.rows[i + 1], 5, r["eventType"], alt=alt)
        _data_cell(table.rows[i + 1], 6, r["eventInvolved"], alt=alt)
        _data_cell(table.rows[i + 1], 7, r["causeOfEvent"], alt=alt)


def _build_section6(doc: Document, data: dict, charts: dict[str, bytes]) -> None:
    _add_heading(doc, "6. Oversikt/Statistikk RUH:")
    cfg = data["config"]
    period_label = cfg.get("periodLabel", f"{cfg.get('quarter', '')} {cfg['year']}")
    _add_body(
        doc,
        f"I {period_label} ble det registrert "
        f"{data['rue']['total']} uoenskede hendelser (RUH).",
    )

    _add_chart_row(doc, charts["rueEventType"], charts["rueEventInvolved"], charts["rueCauseOfEvent"])

    _add_heading(doc, "Hendelsene omfattet", level=2)
    _add_chart(doc, charts["rueEventInvolvedBar"])

    _add_heading(doc, "Rapporteringsfrekvens per maaned", level=2)
    _add_chart(doc, charts["rueMonthlyFrequency"])

    _add_heading(doc, "Hendelsesliste", level=2)
    _add_body(doc, f"Totalt {len(data['rue']['rows'])} hendelser, sortert nyeste foerst.")
    _build_rue_table(doc, data["rue"]["rows"])


def _build_qd_table(doc: Document, rows: list[dict]) -> None:
    headers = ["Avdekket", "Innsendt", "Prosjekt", "Tittel", "Nr.", "Angaar", "I forhold til", "Aarsak"]
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Table Grid"

    for i, h in enumerate(headers):
        _header_cell(table.rows[0], i, h)

    for i, r in enumerate(rows):
        alt = i % 2 == 1
        _data_cell(table.rows[i + 1], 0, _format_date(r.get("eventTime")), alt=alt)
        _data_cell(table.rows[i + 1], 1, _format_date(r["submitDate"]), alt=alt)
        _data_cell(table.rows[i + 1], 2, r["projectName"], alt=alt)
        _data_cell(table.rows[i + 1], 3, r["title"], alt=alt)
        _data_cell(table.rows[i + 1], 4, f"Q-{r['eventId']}", center=True, alt=alt)
        _data_cell(table.rows[i + 1], 5, r["concerning"], alt=alt)
        _data_cell(table.rows[i + 1], 6, r["relatesTo"], alt=alt)
        _data_cell(table.rows[i + 1], 7, r["cause"], alt=alt)


def _build_section7(doc: Document, data: dict, charts: dict[str, bytes]) -> None:
    _add_heading(doc, "7. Oversikt/statistikk Kvalitetsavvik:")
    cfg = data["config"]
    period_label = cfg.get("periodLabel", f"{cfg.get('quarter', '')} {cfg['year']}")
    _add_body(
        doc,
        f"I {period_label} ble det registrert "
        f"{data['qd']['total']} kvalitetsavvik.",
    )

    _add_chart_row(doc, charts["qdConcerning"], charts["qdRelatesTo"], charts["qdCause"])

    _add_heading(doc, "Avviksliste", level=2)
    _add_body(doc, f"Totalt {len(data['qd']['rows'])} avvik, sortert nyeste foerst.")
    _build_qd_table(doc, data["qd"]["rows"])


def _build_section8(doc: Document, data: dict) -> None:
    _add_heading(doc, "8. Samlet oversikt HMSK:")
    cfg = data["config"]
    _add_body(doc, f"Maanedlig oversikt for {cfg['year']}.")

    rows = data["monthlySummary"]
    headers = ["Maaned", "RUH", "RUH-frekv.", "SJA", "Vernerunder", "Kvalitetsavvik", "Sykefravaer %", "Ans."]
    table = doc.add_table(rows=len(rows) + 2, cols=len(headers))
    table.style = "Table Grid"

    for i, h in enumerate(headers):
        _header_cell(table.rows[0], i, h)

    for i, r in enumerate(rows):
        alt = i % 2 == 1
        _data_cell(table.rows[i + 1], 0, r["month"], alt=alt)
        _data_cell(table.rows[i + 1], 1, str(r["ruh"]), center=True, alt=alt)
        _data_cell(table.rows[i + 1], 2, r["ruhFrequency"], center=True, alt=alt)
        _data_cell(table.rows[i + 1], 3, str(r["sja"]), center=True, alt=alt)
        _data_cell(table.rows[i + 1], 4, str(r["vernerunder"]), center=True, alt=alt)
        _data_cell(table.rows[i + 1], 5, str(r["kvalitetsavvik"]), center=True, alt=alt)
        _data_cell(table.rows[i + 1], 6, f"{r['sykefraværProsent']:.1f}%", center=True, alt=alt)
        _data_cell(table.rows[i + 1], 7, str(r["antallAnsatte"]), center=True, alt=alt)

    # Totals row
    total_row = table.rows[-1]
    totals = {
        "ruh": sum(r["ruh"] for r in rows),
        "sja": sum(r["sja"] for r in rows),
        "vernerunder": sum(r["vernerunder"] for r in rows),
        "qa": sum(r["kvalitetsavvik"] for r in rows),
    }
    for col in range(len(headers)):
        _set_cell_shading(total_row.cells[col], HEADER_COLOR)

    _data_cell(total_row, 0, "Totalt", bold=True, color=HEADER_TEXT_COLOR)
    _data_cell(total_row, 1, str(totals["ruh"]), center=True, bold=True, color=HEADER_TEXT_COLOR)
    _data_cell(total_row, 2, "", color=HEADER_TEXT_COLOR)
    _data_cell(total_row, 3, str(totals["sja"]), center=True, bold=True, color=HEADER_TEXT_COLOR)
    _data_cell(total_row, 4, str(totals["vernerunder"]), center=True, bold=True, color=HEADER_TEXT_COLOR)
    _data_cell(total_row, 5, str(totals["qa"]), center=True, bold=True, color=HEADER_TEXT_COLOR)
    _data_cell(total_row, 6, "", color=HEADER_TEXT_COLOR)
    _data_cell(total_row, 7, "", color=HEADER_TEXT_COLOR)


# ---- Main builder ----

def build_document(data: dict, charts: dict[str, bytes]) -> bytes:
    print("\nBygger dokument...")
    config = data["config"]

    doc = Document()
    # Set default font
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(11)

    _build_cover(doc, config)
    _build_toc(doc, config)
    _build_metadata(doc, config)
    _build_section1(doc, config)
    _build_section2(doc, data)
    _build_section3(doc, data)
    _build_section4(doc, config)
    _build_section5(doc, config)
    doc.add_page_break()
    _build_section6(doc, data, charts)
    doc.add_page_break()
    _build_section7(doc, data, charts)
    doc.add_page_break()
    _build_section8(doc, data)

    buf = io.BytesIO()
    doc.save(buf)
    result = buf.getvalue()
    print(f"  Dokument bygget ({len(result) // 1024} KB)")
    return result
